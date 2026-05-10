"""
python make_report.py
-> model_karsilastirma_raporu.docx oluşturur
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Sayfa kenar boşlukları ---
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)

# ---------------------------------------------------------------
# Yardımcı: hücre arka plan rengi
# ---------------------------------------------------------------
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_font(cell, text, bold=False, size=11,
                  color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p   = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

# ---------------------------------------------------------------
# Başlık
# ---------------------------------------------------------------
title = doc.add_heading('Derin Öğrenme Tabanlı Deepfake Tespit Modelleri\nKarşılaştırmalı Değerlendirme Raporu', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)

doc.add_paragraph()

# ---------------------------------------------------------------
# Test seti bilgisi
# ---------------------------------------------------------------
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = info.add_run('Test Seti: ')
r.bold = True
info.add_run('4.601 görüntü  |  Gerçek: 1.510  |  Sahte: 3.091')
info.add_run('\n')
r2 = info.add_run('Kaynak: ')
r2.bold = True
info.add_run('Roop veri seti (%20 test bölümü) + Celeb-DF Pro (eğitimde kullanılmayan kısım)')

doc.add_paragraph()

# ---------------------------------------------------------------
# 1. Tablo — Ana Metrikler
# ---------------------------------------------------------------
h1 = doc.add_heading('1. Model Performans Karşılaştırması', level=2)
for run in h1.runs:
    run.font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)

headers1 = ['Model', 'Accuracy', 'Precision', 'Recall', 'F1-Score', 'AUC-ROC']
rows1 = [
    ['EfficientNet-B4',  '%68.66', '%90.56', '%59.56', '%71.86', '%83.10'],
    ['Xception + FFT',   '%47.27', '%69.28', '%38.66', '%49.63', '%52.63'],
    ['FusionModel',      '%97.96', '%99.12', '%97.83', '%98.47', '%99.80'],
]

tbl1 = doc.add_table(rows=1 + len(rows1), cols=len(headers1))
tbl1.style           = 'Table Grid'
tbl1.alignment       = WD_TABLE_ALIGNMENT.CENTER

# Başlık satırı
for i, h in enumerate(headers1):
    cell = tbl1.rows[0].cells[i]
    set_cell_bg(cell, '1a202c')
    set_cell_font(cell, h, bold=True, size=11, color='FFFFFF')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Veri satırları
row_colors = ['F7FAFC', 'EDF2F7']
fusion_color = 'C6F6D5'   # açık yeşil — en iyi model

for ri, row_data in enumerate(rows1):
    is_fusion = (ri == 2)
    for ci, val in enumerate(row_data):
        cell = tbl1.rows[ri + 1].cells[ci]
        bg = fusion_color if is_fusion else row_colors[ri % 2]
        set_cell_bg(cell, bg)
        align = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
        set_cell_font(cell, val, bold=is_fusion, size=11, align=align)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Sütun genişlikleri
col_widths = [Cm(4.5), Cm(2.4), Cm(2.6), Cm(2.4), Cm(2.4), Cm(2.6)]
for row in tbl1.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_widths[i]

doc.add_paragraph()

# ---------------------------------------------------------------
# 2. Tablo — Karışıklık Matrisi
# ---------------------------------------------------------------
h2 = doc.add_heading('2. Karışıklık Matrisi (Confusion Matrix)', level=2)
for run in h2.runs:
    run.font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)

headers2 = ['Model', 'TP (Doğru Sahte)', 'TN (Doğru Gerçek)', 'FP (Yanlış Alarm)', 'FN (Kaçan Sahte)']
rows2 = [
    ['EfficientNet-B4', '1.841', '1.318', '192',   '1.250'],
    ['Xception + FFT',  '1.195',   '980', '530',   '1.896'],
    ['FusionModel',     '3.024', '1.483',  '27',      '67'],
]

tbl2 = doc.add_table(rows=1 + len(rows2), cols=len(headers2))
tbl2.style     = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(headers2):
    cell = tbl2.rows[0].cells[i]
    set_cell_bg(cell, '1a202c')
    set_cell_font(cell, h, bold=True, size=11, color='FFFFFF')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for ri, row_data in enumerate(rows2):
    is_fusion = (ri == 2)
    for ci, val in enumerate(row_data):
        cell = tbl2.rows[ri + 1].cells[ci]
        # FP ve FN sütunlarını renklendir
        if not is_fusion and ci == 3:   # FP — yanlış alarm
            bg = 'FED7D7'
        elif not is_fusion and ci == 4: # FN — kaçan sahte
            bg = 'FED7D7'
        elif is_fusion:
            bg = fusion_color
        else:
            bg = row_colors[ri % 2]
        set_cell_bg(cell, bg)
        align = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
        set_cell_font(cell, val, bold=is_fusion, size=11, align=align)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

col_widths2 = [Cm(3.8), Cm(3.5), Cm(3.5), Cm(3.5), Cm(3.5)]
for row in tbl2.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_widths2[i]

doc.add_paragraph()

# ---------------------------------------------------------------
# Yorum / Bulgular
# ---------------------------------------------------------------
h3 = doc.add_heading('3. Bulgular ve Yorum', level=2)
for run in h3.runs:
    run.font.color.rgb = RGBColor(0x1a, 0x20, 0x2c)

bullets = [
    ('EfficientNet-B4: ', 'Piksel tabanlı özellik çıkarımında yüksek precision (%90.56) elde etmiş; '
     'ancak sahte görüntülerin %40\'ını (FN=1.250) kaçırmaktadır. Tek başına yetersizdir.'),
    ('Xception + FFT: ', 'Frekans domeninde çalışarak farklı bir perspektif sunmaktadır; '
     'fakat bu test setinde düşük recall (%38.66) ve AUC-ROC (%52.63) ile sınırlı kalmıştır.'),
    ('FusionModel: ', 'Her iki modelin güçlü yönlerini özellik düzeyinde birleştirerek '
     'dramatik bir iyileşme sağlamıştır. Accuracy %97.96, F1-Score %98.47, AUC-ROC %99.80 '
     'ile state-of-the-art seviyesine ulaşılmıştır. Sadece 27 yanlış alarm ve 67 kaçan sahte '
     'görüntü ile oldukça güvenilir bir sistem elde edilmiştir.'),
]

for bold_part, normal_part in bullets:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(bold_part)
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(normal_part)
    r2.font.size = Pt(11)

doc.add_paragraph()

# ---------------------------------------------------------------
# Kaydet
# ---------------------------------------------------------------
out_path = 'model_karsilastirma_raporu.docx'
doc.save(out_path)
print(f"Rapor kaydedildi: {out_path}")
